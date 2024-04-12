import tkinter as tk
from tkinter import filedialog, messagebox, font
from tkinter.scrolledtext import ScrolledText
import docx
from docx2pdf import convert
from PIL import Image, ImageTk
import io
import base64
import os


logo_base64 = """/9j/4AAQSkZJRgABAQEAYABgAAD/4REiRXhpZgAATU0AKgAAAAgABAE7AAIAAAAdAAAISodpAAQAAAABAAAIaJydAAEAAAA6AAAQ4OocAAcAAAgMAAAAPgAAAAAc6gAAAAgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAEplcnJ5IENyaXN0b3BoZXIgTHVjYXMgUG9uY2UAAAAFkAMAAgAAABQAABC2kAQAAgAAABQAABDKkpEAAgAAAAMyNwAAkpIAAgAAAAMyNwAA6hwABwAACAwAAAiqAAAAABzqAAAACAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAMjAyMzowNToxMCAxNTowMjozMQAyMDIzOjA1OjEwIDE1OjAyOjMxAAAASgBlAHIAcgB5ACAAQwByAGkAcwB0AG8AcABoAGUAcgAgAEwAdQBjAGEAcwAgAFAAbwBuAGMAZQAAAP/hCy9odHRwOi8vbnMuYWRvYmUuY29tL3hhcC8xLjAvADw/eHBhY2tldCBiZWdpbj0n77u/JyBpZD0nVzVNME1wQ2VoaUh6cmVTek5UY3prYzlkJz8+DQo8eDp4bXBtZXRhIHhtbG5zOng9ImFkb2JlOm5zOm1ldGEvIj48cmRmOlJERiB4bWxuczpyZGY9Imh0dHA6Ly93d3cudzMub3JnLzE5OTkvMDIvMjItcmRmLXN5bnRheC1ucyMiPjxyZGY6RGVzY3JpcHRpb24gcmRmOmFib3V0PSJ1dWlkOmZhZjViZGQ1LWJhM2QtMTFkYS1hZDMxLWQzM2Q3NTE4MmYxYiIgeG1sbnM6ZGM9Imh0dHA6Ly9wdXJsLm9yZy9kYy9lbGVtZW50cy8xLjEvIi8+PHJkZjpEZXNjcmlwdGlvbiByZGY6YWJvdXQ9InV1aWQ6ZmFmNWJkZDUtYmEzZC0xMWRhLWFkMzEtZDMzZDc1MTgyZjFiIiB4bWxuczp4bXA9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC8iPjx4bXA6Q3JlYXRlRGF0ZT4yMDIzLTA1LTEwVDE1OjAyOjMxLjI3MjwveG1wOkNyZWF0ZURhdGU+PC9yZGY6RGVzY3JpcHRpb24+PHJkZjpEZXNjcmlwdGlvbiByZGY6YWJvdXQ9InV1aWQ6ZmFmNWJkZDUtYmEzZC0xMWRhLWFkMzEtZDMzZDc1MTgyZjFiIiB4bWxuczpkYz0iaHR0cDovL3B1cmwub3JnL2RjL2VsZW1lbnRzLzEuMS8iPjxkYzpjcmVhdG9yPjxyZGY6U2VxIHhtbG5zOnJkZj0iaHR0cDovL3d3dy53My5vcmcvMTk5OS8wMi8yMi1yZGYtc3ludGF4LW5zIyI+PHJkZjpsaT5KZXJyeSBDcmlzdG9waGVyIEx1Y2FzIFBvbmNlPC9yZGY6bGk+PC9yZGY6U2VxPg0KCQkJPC9kYzpjcmVhdG9yPjwvcmRmOkRlc2NyaXB0aW9uPjwvcmRmOlJERj48L3g6eG1wbWV0YT4NCiAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAKICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgIAogICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgCiAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAKICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgIAogICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgCiAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAKICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgIAogICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgCiAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAKICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgIAogICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgCiAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAKICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgIAogICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgCiAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAKICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgIAogICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgCiAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAKICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgIAogICAgICAgICAgICAgICAgICAgICAgICAgICAgPD94cGFja2V0IGVuZD0ndyc/Pv/bAEMABwUFBgUEBwYFBggHBwgKEQsKCQkKFQ8QDBEYFRoZGBUYFxseJyEbHSUdFxgiLiIlKCkrLCsaIC8zLyoyJyorKv/bAEMBBwgICgkKFAsLFCocGBwqKioqKioqKioqKioqKioqKioqKioqKioqKioqKioqKioqKioqKioqKioqKioqKioqKv/AABEIAJkAzgMBIgACEQEDEQH/xAAfAAABBQEBAQEBAQAAAAAAAAAAAQIDBAUGBwgJCgv/xAC1EAACAQMDAgQDBQUEBAAAAX0BAgMABBEFEiExQQYTUWEHInEUMoGRoQgjQrHBFVLR8CQzYnKCCQoWFxgZGiUmJygpKjQ1Njc4OTpDREVGR0hJSlNUVVZXWFlaY2RlZmdoaWpzdHV2d3h5eoOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4eLj5OXm5+jp6vHy8/T19vf4+fr/xAAfAQADAQEBAQEBAQEBAAAAAAAAAQIDBAUGBwgJCgv/xAC1EQACAQIEBAMEBwUEBAABAncAAQIDEQQFITEGEkFRB2FxEyIygQgUQpGhscEJIzNS8BVictEKFiQ04SXxFxgZGiYnKCkqNTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqCg4SFhoeIiYqSk5SVlpeYmZqio6Slpqeoqaqys7S1tre4ubrCw8TFxsfIycrS09TV1tfY2dri4+Tl5ufo6ery8/T19vf4+fr/2gAMAwEAAhEDEQA/APpGiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigDP1vXLLw/p5vNSkZY87VCrlnbBOAPopOTgAAkkAVyQ+KUU5t/suhXuLiJp0+0SxRFYR1lf5jsT0LYz2zS/FKLzrXRl+zQXOLxm2XMmyEbYnO+Q/3FxuI74x3rhP+Pj/p7+1f6T/pf7v7dt/5ern/AJ52yfwR/wAWBWsIJq7M5SaZ3Nv8U45fJMuhXqieJ502Sx58lesrB2XYh7FsZ7Zrq9B8QWPiTTRe6a0nl8BkljKOhKhgCD7MCCOCCCDXjf8Ax8f9Pf2r/Sf9L/d/btv/AC9XP/PO2T+CP+LArE0fxJqemfF3T1sNRvZLfUrq3SdrkbRdq5VTJ5fRVIPyDsu2qdNW0EpPqfSVFFFYGpja54p07QHWK6Mktw0bSiCFQWEa9XYkhVUerECufi+KlhN5GzQtYH2iBrlN4t1xEvWRszDYvozYB7Zrk/E8sl3421mKRoZwJ0MdnI2EKxxIfNuW7QISWC/xMx9Kzf8Aj4/6e/tX+k/6X+7+3bf+Xq5/552yfwR/xYFbxgramTk7nfRfFSwm+z7NC1gfaIGuU3i3XES9ZGzMNi+jNgHtmiL4qWE32fZoWsD7RA1ym8W64iXrI2ZhsX0ZsA9s1wP/AB8f9Pf2r/Sf9L/d/btv/L1c/wDPO2T+CP8AiwKP+Pj/AKe/tX+k/wCl/u/t23/l6uf+edsn8Ef8WBT9nEOdnfRfFSwm+z7NC1gfaIGuU3i3XES9ZGzMNi+jNgHtmnwfFPSZREX03VYRNA1ypkgTiEdZGAclV9CcA9q8+/4+P+nv7V/pP+l/u/t23/l6uf8AnnbJ/BH/ABYFH/Hx/wBPf2r/AEn/AEv939u2/wDL1c/887ZP4I/4sCj2cQ52ep2nxA8MXbRr/aiWzyoJEW8Rrfcp6EFwAQexHBro0dZEDxsGVhkMDkEV4T/x8f8AT59q/wBJ/wBM+T7dt/5ern/nnbJ/BH/FgU7Sbm80aZJfDt5JbvMjTIsvyR3S7svdTxn5YYByECgMexqXT7DU+57rRXLeEvG1t4iCWt0n2TUTF5yxMCq3EWSBNHu52nGcHkfTBPU1k01uWncKKKKQwooooAKKKKACiiigAooooA4P4pReda6Mv2aC5xeM2y5k2QjETnfIf7i43Ed9uO9cJ/x8f9Pf2r/Sf9L/AHf27b/y9XP/ADztk/gj/iwK7v4pReda6Mv2aC5xeM2y5k2QjETnfIf7i43Ed9uO9cJ/x8f9Pf2r/Sf9L/d/btv/AC9XP/PO2T+CP+LArop/CZS3D/j4/wCnv7V/pP8Apf7v7dt/5ern/nnbJ/BH/FgVySP53xl0Sbzrm483UbN/tNyuxp8yL84X+FT/AAjsu2ut/wCPj/p7+1f6T/pf7v7dt/5ern/nnbJ/BH/FgVySP53xl0Sbzrm483UbN/tNyuxp8yL84X+FT/COy7avoStz6kooorkNzxbxCPM8Za3Btt5t12JfsrHar7IkPm3T9oI+oX+Jiao/8fH/AE9/av8ASf8AS/3f27b/AMvVz/zztk/gj/iwKveIR5njLW4NtvNuuxL9lY7VfZEh826ftBH1C/xMTVH/AI+P+nv7V/pP+l/u/t23/l6uf+edsn8Ef8WBXVHZGD3GSyBoXlYNeidTdkXKlTfbP+Xq4A+5bJj5Iv4sDisKXxro7faN7zXm6RXf7TGR9vm7ST46Qp/DEuc45rX1H/SNLuv9dd/aoXuf3n7uS+2qf9Kn/wCedun/ACzj/iIFQfBXwdoPiuPWT4gsBdm2MHlZldNu7fn7pGfujrTbsrglcoP4z0eRpw8sl3vkV3N1GQL6btJPjOIU/hiXOcc1o22tadqc0kVveRXzyzbmF23l/bZFGTPcHolvGPuxDrjpXp5+D3gRhzoQ/C6mH/s9Ur/4H+DbuPFtbXVi3ZoLhj/6HuqPaRK5GcX/AMfH/T39q/0n/S/3f27b/wAvVz/zztk/gj/iwKP+Pj/p7+1f6T/pf7v7dt/5ern/AJ52yfwR/wAWBVjUvDl94b1STSb6YarHMou4ZrhWVbkLwWu5CT+7hAU7Bw29ar/8fH/T39q/0n/S/wB39u2/8vVz/wA87ZP4I/4sCrTuRsRTRS3RSazuplvJCbu3vHIimmZQQbyZiP3NuoJCR/xA+9el/DfxwnjXQGecCPUrMiO7jAwCT0cezYPHYgj0rzn/AI+P+nv7V/pP+l/u/t23/l6uf+edsn8Ef8WBWf4R1eTQfjPaSm4uJLXW1CPNcIEa538LIE/hUyDKjrt+tTON0VF2Z9D0UUVzGwUUUUAFFFFABRRRQAUUUUAcH8UovOtdGX7NBc4vGbZcybIRiJzvkP8AcXG4jvtx3rhP+Pj/AKe/tX+k/wCl/u/t23/l6uf+edsn8Ef8WBXd/FKLzrXRl+zQXOLxm2XMmyEYic75D/cXG4jvtx3rhP8Aj4/6e/tX+k/6X+7+3bf+Xq5/552yfwR/xYFdFP4TKW4f8fH/AE9/av8ASf8AS/3f27b/AMvVz/zztk/gj/iwK5JH874y6JN51zcebqNm/wBpuV2NPmRfnC/wqf4R2XbXW/8AHx/09/av9J/0v939u2/8vVz/AM87ZP4I/wCLArkkfzvjLok3nXNx5uo2b/abldjT5kX5wv8ACp/hHZdtX0JW59SUUUVyG54t4hHmeMtbg228267Ev2VjtV9kSHzbp+0EfUL/ABMTVH/j4/6e/tX+k/6X+7+3bf8Al6uf+edsn8Ef8WBV7xCPM8Za3Btt5t12JfsrHar7IkPm3T9oI+oX+Jiao/8AHx/09/av9J/0v939u2/8vVz/AM87ZP4I/wCLArqjsjB7lbUf9I0u6/1139qhe5/efu5L7ap/0qf/AJ526f8ALOP+IgVqfs5/6rxD/vW3/tWsvUf9I0u6/wBdd/aoXuf3n7uS+2qf9Kn/AOedun/LOP8AiIFan7Of+q8Q/wC9bf8AtWlP4WOO57ZRRR061zGx5/8AFK2EzaKwgiuHWWUGO4k2Qbdm4tKe8alFYjvtAriP+Pj/AKe/tX+k/wCl/u/t23/l6uf+edsn8Ef8WBW3411a38R+IoFtvLu7eyjf7PFI22CU5HmXEzf88EKKB/fYEDjrif8AHx/09/av9J/0v939u2/8vVz/AM87ZP4I/wCLArpgrRMZbh/x8f8AT39q/wBJ/wBL/d/btv8Ay9XP/PO2T+CP+LArjtbufK+IekX0f2mZmeCb7VccPdnzD+9CfwKcYVf7qqe9dj/x8f8AT39q/wBJ/wBL/d/btv8Ay9XP/PO2T+CP+LArirAf8JF8XbEQXM1+s1/FuuZVw0gUjc+3+FeCQvYYHaqJPqyiiiuQ6AooooAKKKKACiiigAooooA4P4pReda6Mv2aC5xeM2y5k2QjETnfIf7i43Ed9uO9cJ/x8f8AT39q/wBJ/wBL/d/btv8Ay9XP/PO2T+CP+LAru/ilF51roy/ZoLnF4zbLmTZCMROd8h/uLjcR32471wn/AB8f9Pf2r/Sf9L/d/btv/L1c/wDPO2T+CP8AiwK6KfwmUtw/4+P+nv7V/pP+l/u/t23/AJern/nnbJ/BH/FgVySP53xl0Sbzrm483UbN/tNyuxp8yL84X+FT/COy7a63/j4/6e/tX+k/6X+7+3bf+Xq5/wCedsn8Ef8AFgVySP53xl0Sbzrm483UbN/tNyuxp8yL84X+FT/COy7avoStz6kooorkNzxbxCPM8Za3Btt5t12JfsrHar7IkPm3T9oI+oX+Jiao/wDHx/09/av9J/0v939u2/8AL1c/887ZP4I/4sCr3iEeZ4y1uDbbzbrsS/ZWO1X2RIfNun7QR9Qv8TE1R/4+P+nv7V/pP+l/u/t23/l6uf8AnnbJ/BH/ABYFdUdkYPciuYf7QtZYd7z/AG5DPvuD5TXeBj7XcH/llbJ/BH3wOKZ8O7jVPA63Qs47O8/tQpJH57tFtgj3ZuH4+SL5uC2CccCrH/Hx/wBPf2r/AEn/AEv939u2/wDL1c/887ZP4I/4sCj/AI+P+nv7V/pP+l/u/t23/l6uf+edsn8Ef8WBTaurMS0OnPxJ1+eOP7Lp+mo8+ZIfOeQDyB1uHzjy4vQnluwrI1DxFrevOVvbxJbWaMvbWSxm2imjH3rm45LLAOylsvgcDOKz/wDj4/6e/tX+k/6X+7+3bf8Al6uf+edsn8Ef8WBR/wAfH/T39q/0n/S/3f27b/y9XP8Azztk/gj/AIsCkopD5mNRFeNUjVblLgC4VboeWLwJ/wAvNwP+WdqmPkj74HFO/wCPj/p7+1f6T/pf7v7dt/5ern/nnbJ/BH/FgVFdXCpayzukl+JlN2UmQhr/AGf8vM6jlLZP4Iv4uK4vXfGsl3ut9N8xo5XEt1cXKgyXsg6F1GQEX+GPlR3zVCNfxXrsdtpskKTyTyXpErPKNsl4e00o/giH/LOL6MeMA7/wF8JyXGqz+J7uPEFurQ2pYfekI+Zh7BSR/wAC9jWH4I+FOteMb0alrpns9OkbzHnmz5txnn5Qeef7x45719G6fp9ppWnQWOnQLb20CBI40HCj+v171lOWlkXGPUs0UUVgahRRRQAUUUUAFFFFABRRRQBwfxSi8610Zfs0Fzi8ZtlzJshGInO+Q/3FxuI77cd64T/j4/6e/tX+k/6X+7+3bf8Al6uf+edsn8Ef8WBXd/FKLzrXRl+zQXOLxm2XMmyEYic75D/cXG4jvtx3rhP+Pj/p7+1f6T/pf7v7dt/5ern/AJ52yfwR/wAWBXRT+EyluH/Hx/09/av9J/0v939u2/8AL1c/887ZP4I/4sCuSR/O+MuiTedc3Hm6jZv9puV2NPmRfnC/wqf4R2XbXW/8fH/T39q/0n/S/wB39u2/8vVz/wA87ZP4I/4sCuSR/O+MuiTedc3Hm6jZv9puV2NPmRfnC/wqf4R2XbV9CVufUlFFFchueLeIR5njLW4NtvNuuxL9lY7VfZEh826ftBH1C/xMTVH/AI+P+nv7V/pP+l/u/t23/l6uf+edsn8Ef8WBV7xCPM8Za3Btt5t12JfsrHar7IkPm3T9oI+oX+Jiao/8fH/T39q/0n/S/wB39u2/8vVz/wA87ZP4I/4sCuqOyMHuRXE3+iyzbHv/ADkN5suBtbUNgz9puB/yztlx8kf8XFcePH0sn/HxpwuvMbzrnzZCftcw+4ZcDmNP4Yhgeua6zUf9I0u6/wBdd/aoXuf3n7uS+2qf9Kn/AOedun/LOP8AiIFaX7OqI8XiHeqthrfGRn/npRJ2VwSuzjrPxvbXWU1S2bfMyyTNcNuju584DzsACIUHKxKCOK6cEXIyD9sF1/pObv5Pt23/AJern/nnbJ/BH/FgV7XeaNpmoRmO/wBOtLlG6rNArg/mK8u8WeD7fwtfRnSbeP8Ase/k3NazOEgjuFGVMzn5jCFDvszjcuO4FTGabsU4tGN/x8f9Pf2r/Sf9L/d/btv/AC9XP/PO2T+CP+LAqC2Eena5DrtrYw6ncsrTst7EvmXsYOZLptx22yDGIzjJ75qf/j4/6e/tX+k/6X+7+3bf+Xq5/wCedsn8Ef8AFgUf8fH/AE9/av8ASf8AS/3f27b/AMvVz/zztk/gj/iwKsg9l0DX9P8AEujQ6npM3mwSjoeGRu6sOxH+eK0q8G8N+Jf+EL8bW8808j6Rr5/fyXBCO75/4+vLAHlozMQAeqgnsK95rmlHlZtF3QUUUVJQUUUUAFFFFABRRVLWtSTRtDvdSkXetpA82zON20Zxn36UAXaK8z0z4l6qs6HVdPt7q2nVpojaAwyrAoy0zJIxUR5yFZmUng45rfT4l+HHiDeZeq5g+0iNrCYHyv8Anpnbt2/7Wce9U4yRPMmZvxSAmTRbcW0N0/2iSby7l9sKqkZy8p/uKWViO+Md64f/AI+P+nv7V/pP+l/u/t23/l6uf+edsn8Ef8WBV7xBrMni3WoNQktttstuTp+n3TbFeLIZrq5x92EFVIXPzlF/Gj/x8f8AT39q/wBJ/wBL/d/btv8Ay9XP/PO2T+CP+LAreKsjOTuw/wCPj/p7+1f6T/pf7v7dt/5ern/nnbJ/BH/FgVyMUv2j4zaNKs1zcmTUbNzcXKbGny6HeF/hU/wjsu2uu/4+P+nv7V/pP+l/u/t23/l6uf8AnnbJ/BH/ABYFcn4eZta+NmnSQ3Ut+f7QjkNxKu0yiPBLbf4R8vC9hgdqp7CW59R0UUVyG54t4hHmeMtbg228267Ev2VjtV9kSHzbp+0EfUL/ABMTVH/j4/6e/tX+k/6X+7+3bf8Al6uf+edsn8Ef8WBV7xCPM8Za3Btt5t12JfsrHar7IkPm3T9oI+oX+Jiao/8AHx/09/av9J/0v939u2/8vVz/AM87ZP4I/wCLArqjsjB7lbUf9I0u6/1139qhe5/efu5L7ap/0qf/AJ526f8ALOP+IgVqfs5/6rxD/vW3/tWsvUf9I0u6/wBdd/aoXuf3n7uS+2qf9Kn/AOedun/LOP8AiIFan7Of+q8Q/wC9bf8AtWlP4WOO57ZWZ4j0aPxB4eu9NlCZmT92zruCuOUYjuAwBx3rTormNjwSBpbqLF1F5lzMTLPBduAbiWM7XmuiOEt4mBVY/wCIrUn/AB8f9Pf2r/Sf9L/d/btv/L1c/wDPO2T+CP8AiwK2vGekf2b4uuUit45LbVMXywuoSF5UADtcSdfKjwJNv8TSVi/8fH/T39q/0n/S/wB39u2/8vVz/wA87ZP4I/4sCupO6uYNWZn65YHWNJmhjklmluR9rjlkCxy3ZUYN1MTxDbqu5Y04zmvU/hT4m/4SbwHavO+67s/9FuMnklR8re+Vxz65rzz/AI+P+nv7V/pP+l/u/t23/l6uf+edsn8Ef8WBR8PdX/4Rr4otaSyytp/iKMPFNOixGaTJ2yBB91WfeFHoy1M1dDi7M93ooornNgooooAKKKKACud8fkjwFquDGP3IH737n3h9729a6Ks7xBpn9teHNQ01Sqtc27xozjKhiPlJHpnFNbiex4z/AMfH/T39q/0n/S/3f27b/wAvVz/zztk/gj/iwKP+Pj/p7+1f6T/pf7v7dt/5ern/AJ52yfwR/wAWBUM87W98+n6xH5WoSSK81vqLBPtkoXJnnf7v2aPB2RqTuxU3/Hx/09/av9J/0v8Ad/btv/L1c/8APO2T+CP+LArqMA/4+P8Ap7+1f6T/AKX+7+3bf+Xq5/552yfwR/xYFH/Hx/09/av9J/0v939u2/8AL1c/887ZP4I/4sCj/j4/6e/tX+k/6X+7+3bf+Xq5/wCedsn8Ef8AFgVl6n4i02xspLie4GoGaTeIZf8AWajKOk06j7kC/wAEXGeD05oAm1jVrfTdNe8uy94LkieOOdcPqTjhZ5l/gt1P+ri/ix6VD8C9Jl1Xx/PrE5ZxYQtI0hOcySZUZ+oLn8K4u5u9V8Y62lnZpNd3N1LnbxvmfGNzdgAM4H3UX8Sfpb4e+DYvBXhaKxOx7yU+bdyr/E/oD6AcD8T3qJuysVFXZ1NFFFc5seLeIR5njLW4NtvNuuxL9lY7VfZEh826ftBH1C/xMTVH/j4/6e/tX+k/6X+7+3bf+Xq5/wCedsn8Ef8AFgVf8Yp/Z3jDUYtRW3SK+uklto7gGOG6IjX555TgGKMgnywclj05FZf2q3uP+W8d59qfzv8AS3CfbmX/AJebn/nnbJ/BF/FgcV1R2Rg9yPUf9I0u6/1139qhe5/efu5L7ap/0qf/AJ526f8ALOP+IgVqfs5/6rxD/vW3/tWue1vWdPj0e5ee7N39ryyqzBZb+UDAnmA/1cKf8s4u+AeldD+zn/qvEP8AvW3/ALVpT+Fjjue2UUUVzGxyPxK0n+0PCb3kcEU8+mOLtY5gSjqv31YDkrj5tvcqK82/4+P+nv7V/pP+l/J9u2/8vVz/AM87ZP4I/wCLAr3ZlV0KuAysMEEZBFeFapp6+HNUm0a9WIKZ82yXcjCO9VeY5Z5W48mJNq+WpOWU8c1tTfQzmuon/Hx/09/av9J/0v8Ad/btv/L1c/8APO2T+CP+LArE8TWklzZJqNlJNJexE30VwYsXF0q43XL/APPGEBQI171rfaYLj/lvHd/an87/AEtwn25l/wCXm5/552yfwRfxYHFRT6hYvGfPu4bhbx/M/wBMkEZv3H/Le4/5526/wRDluOK1Mz2nwnr0fibwrYavFgG4iBkUfwuOGH4MDWxXivwA10eZq/h7zfMjQ/a7c7cZXIRj7Z+Tj617VXNJWdjdO6CiiipGFFFFABRRRQBzXjDwFonjWCMatFIk8IxFcwNtkUenIII9iPXGM15hqHwC1eBZE0PxFFJFIRujuVeIEDoDt3Zx9K90oqlJrYTimfPUnwN8a3UspudWsHMhBkeS5lO/HTPyZOPetLSf2ebhnDa7rkaqDzHZxliw/wB5sY/I17nRVe0kLlRz/hfwRoXhC28vRrNVmIw9zJ80sn1bsOOgwPaugoorPcoKKKKAMrxJ4dsPFWhzaVqqM0EuDuRsMjDowPqK8quf2doWkY2fiSRIz0SS0DH8w4/lXtVFUpNbCaTPEov2dF3gz+JSVzyEsuT+Jf8ApXpXgzwRpfgjTZbXS/Mkedg008py0hHTpwAMnA966Oihyb3BRSCiiipGFY/ijwxp3i7Q5NL1ZGMTEOjocNG46Mp9eT+dbFFAHic37OqmRjb+JmCE8K9lkgfUP/SmD9nRifn8T4+ljnP/AJEr2+ir9pInlRw/gL4Xad4Fupr2G8mvb2aLyTI6hFVcgkBRnqQOST0ruKKKltvcdrBRRRSGFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFAH//2Q=="""

app = tk.Tk()
app.title("GESTOR DE PLANTILLAS JR")
app.config(bg='#D6EAF8')

saved_file_path = None  # Agrega esta línea al principio del código

def browse_file():
    global file_path
    file_path = filedialog.askopenfilename(filetypes=[("Documento Word", "*.docx")])
    file_label.config(text=file_path)

def open_and_edit():
        global document_text, file_path
        try:
            document = docx.Document(file_path)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            document_text.delete(1.0, tk.END)
            document_text.insert(tk.END, '\n'.join(full_text))
        except Exception as e:
            messagebox.showerror("Error", f"Error al abrir el documento: {str(e)}")

def open_and_edit_second_file():
    global second_file_path
    second_file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if second_file_path:
        file_name = os.path.basename(second_file_path)
        second_file_label.config(text=file_name)

def save_file():
    global saved_file_path, file_path, second_file_path
    try:
        # Leer el contenido del primer documento
        source_document = docx.Document(file_path)
        source_text = []
        for para in source_document.paragraphs:
            source_text.append(para.text)

        # Cargar el segundo documento
        document = docx.Document(second_file_path)

        # Agregar campos al inicio del segundo documento
        first_paragraph = document.paragraphs[0]
        empresa_paragraph = first_paragraph.insert_paragraph_before(f"Empresa: {entry_empresa.get()}")
        motivo_estudio_paragraph = empresa_paragraph.insert_paragraph_before(f"Motivo del estudio: {entry_motivo_estudio.get()}")
        tipo_estudio_paragraph = motivo_estudio_paragraph.insert_paragraph_before(f"Tipo de estudio: {entry_tipo_estudio.get()}")
        nombre_paragraph = tipo_estudio_paragraph.insert_paragraph_before(f"Nombre: {entry_nombre.get()}")
        cedula_paragraph = nombre_paragraph.insert_paragraph_before(f"Cédula: {entry_cedula.get()}")
        fecha_paragraph = cedula_paragraph.insert_paragraph_before(f"Fecha: {entry_fecha.get()}")

        # Agregar el contenido del primer archivo al segundo archivo
        for para in source_text:
            document.add_paragraph(para)

        saved_file_path = second_file_path

        document.save(saved_file_path)

        print(f"Archivo guardado en: {saved_file_path}")
        messagebox.showinfo("Éxito", "El contenido del primer archivo se copió correctamente al segundo archivo y se agregaron los campos.")
    except Exception as e:
        messagebox.showerror("Error", f"Error al copiar el contenido del primer archivo al segundo archivo y agregar los campos: {str(e)}")


def convert_to_pdf(saved_file_path):
    try:
        if not os.path.exists(saved_file_path):
            raise FileNotFoundError("El archivo de Word no existe en la ruta especificada.")

        print(f"Convirtiendo el archivo: {saved_file_path}")
        output_pdf_path = saved_file_path.replace(".docx", ".pdf")

        # Utiliza docx2pdf para convertir el archivo
        convert(saved_file_path, output_pdf_path)

        messagebox.showinfo("Éxito", "El archivo se convirtió a PDF correctamente.")
    except Exception as e:
        messagebox.showerror("Error", f"Error al convertir el archivo a PDF: {str(e)}")

def browse_file():
    global file_path
    file_path = filedialog.askopenfilename(filetypes=[("Documento Word", "*.docx")])
    file_label.config(text=file_path)

def save_file():
    global saved_file_path, second_file_path
    try:
        # Cargar el segundo documento en lugar del primero
        document = docx.Document(second_file_path)
        first_paragraph = document.paragraphs[0]

        empresa_paragraph = first_paragraph.insert_paragraph_before(f"Empresa: {entry_empresa.get()}")
        motivo_estudio_paragraph = empresa_paragraph.insert_paragraph_before(f"Motivo del estudio: {entry_motivo_estudio.get()}")
        tipo_estudio_paragraph = first_paragraph.insert_paragraph_before(f"Tipo de estudio: {entry_tipo_estudio.get()}")
        nombre_paragraph = tipo_estudio_paragraph.insert_paragraph_before(f"Nombre: {entry_nombre.get()}")
        cedula_paragraph = nombre_paragraph.insert_paragraph_before(f"Cédula: {entry_cedula.get()}")

        # Agregar el contenido del ScrolledText al documento
        document_text_content = document_text.get(1.0, tk.END).strip()
        document.add_paragraph(document_text_content)

        cedula = entry_cedula.get().replace(" ", "_")
        nombre = entry_nombre.get().replace(" ", "_")
        tipo_estudio = entry_tipo_estudio.get().replace(" ", "_")
        fecha = entry_fecha.get().replace(" ", "_")
        saved_file_name = f"{cedula}-{nombre}-{tipo_estudio}-{fecha}jp.docx"
        saved_file_path = os.path.join(os.path.dirname(second_file_path), saved_file_name)

        document.save(saved_file_path)

        print(f"Archivo guardado en: {saved_file_path}")
        messagebox.showinfo("Éxito", "El archivo se guardó correctamente.")
    except Exception as e:
        messagebox.showerror("Error", f"Error al guardar el archivo: {str(e)}")


def convert_to_pdf():
    global saved_file_path  # Agrega esta línea
    try:
        if not os.path.exists(saved_file_path):
            raise FileNotFoundError("El archivo de Word no existe en la ruta especificada.")

        print(f"Convirtiendo el archivo: {saved_file_path}")
        output_pdf_path = saved_file_path.replace(".docx", ".pdf")

        # Utiliza docx2pdf para convertir el archivo
        convert(saved_file_path, output_pdf_path)

        messagebox.showinfo("Éxito", "El archivo se convirtió a PDF correctamente.")
    except Exception as e:
        messagebox.showerror("Error", f"Error al convertir el archivo a PDF: {str(e)}")




file_path = ""
saved_file_path = ""

logo_image_data = io.BytesIO(base64.b64decode(logo_base64))
logo_image = Image.open(logo_image_data)
logo_image = logo_image.resize((150, 110), Image.LANCZOS)
logo_image = ImageTk.PhotoImage(logo_image)

logo_label = tk.Label(app, image=logo_image)
logo_label.grid(row=0, column=2, rowspan=6, padx=(0,10))

label_font = ('Helvetica', 10, 'bold')
label_color = "#D6EAF8"  # Azul oscuro
text_color = "#000000"  # Blanco
entry_bg_color = "#E0BBE4"
entry_font = ('Helvetica', 10, 'bold')  # Fuente Helvetica, tamaño 10 y negrita
button_bg_color = "#89CFF0"
app_bg_color = "#D6EAF8"

tk.Label(app, text="Nombre:", font=label_font, bg=label_color, fg=text_color).grid(row=0, column=0)
entry_nombre = tk.Entry(app, bg=entry_bg_color, font=entry_font)
entry_nombre.grid(row=0, column=1)

tk.Label(app, text="Cédula:", font=label_font, bg=label_color, fg=text_color).grid(row=1, column=0)
entry_cedula = tk.Entry(app, bg=entry_bg_color, font=entry_font)
entry_cedula.grid(row=1, column=1)

tk.Label(app, text="Fecha:", font=label_font, bg=label_color, fg=text_color).grid(row=2, column=0)
entry_fecha = tk.Entry(app, bg=entry_bg_color, font=entry_font)
entry_fecha.grid(row=2, column=1)

tk.Label(app, text="Tipo de estudio:", font=label_font, bg=label_color, fg=text_color).grid(row=3, column=0)
entry_tipo_estudio = tk.Entry(app, bg=entry_bg_color, font=entry_font)
entry_tipo_estudio.grid(row=3, column=1)

tk.Label(app, text="Motivo del estudio:", font=label_font, bg=label_color, fg=text_color).grid(row=4, column=0)
entry_motivo_estudio = tk.Entry(app, bg=entry_bg_color, font=entry_font)
entry_motivo_estudio.grid(row=4, column=1)

tk.Label(app, text="Empresa: ", font=label_font, bg=label_color, fg=text_color).grid(row=5, column=0)
entry_empresa = tk.Entry(app, bg=entry_bg_color, font=entry_font)
entry_empresa.grid(row=5, column=1)

# Agrega el widget ScrolledText aquí
document_text = ScrolledText(app, wrap=tk.WORD, width=50, height=15)
document_text.grid(row=9, column=0, columnspan=3, padx=10, pady=10)

browse_button = tk.Button(app, text="Selecc.. Archivo", command=browse_file, bg=button_bg_color)
browse_button.grid(row=6, column=0, pady=(5,1))

file_label = tk.Label(app, text="", bg=app_bg_color)
file_label.grid(row=6, column=1)

# Agrega el botón open_and_edit_button aquí
copy_content_button = tk.Button(app, text="Guardar Archivo", command=save_file, bg=button_bg_color)
copy_content_button.grid(row=7, column=0, pady=(3,3))

open_and_edit_second_button = tk.Button(app, text="Escoger Membrete", command=open_and_edit_second_file, bg=button_bg_color)
open_and_edit_second_button.grid(row=6, column=2, pady=(5,1))

second_file_label = tk.Label(app, text="", bg=app_bg_color)
second_file_label.grid(row=7, column=2)

open_and_edit_button = tk.Button(app, text="Mostrar Contenido", command=open_and_edit, bg=button_bg_color)
open_and_edit_button.grid(row=8, column=1)

convert_button = tk.Button(app, text="Convertir a PDF", command=convert_to_pdf, bg=button_bg_color)
convert_button.grid(row=8, columnspan=1)

app.mainloop()
